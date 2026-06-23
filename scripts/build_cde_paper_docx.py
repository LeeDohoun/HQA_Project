from __future__ import annotations

import argparse
import re
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "cde_2026_refs" / "2026_summer_paper_template.docx"
SOURCE_MD = ROOT / "docs" / "cde_2026_paper_draft_ko.md"
OUT_DOCX = ROOT / "docs" / "cde_2026_paper_draft_ko.docx"
FIG_DIR = ROOT / "artifacts" / "paper_figures"
FIG_PATH = FIG_DIR / "temporal_rag_architecture.png"
AUDIT_JSON = ROOT / "artifacts" / "paper_temporal_audit" / "temporal_rag_leakage_audit.json"

PROTOCOL_ROWS = [
    ("Theme scope", "AI industry theme candidate selection"),
    ("Periods", "Validation: 2023, 2024, 2026Q1; tuning/reference: 2025"),
    ("Horizons", "Short: 5 trading days, Long: 60 trading days"),
    (
        "Comparators",
        "Deterministic, RSI, Bollinger Band, momentum, LLM-only, hybrid multi-agent",
    ),
    ("Cost assumptions", "Transaction cost 15bp, slippage 5bp, market impact 5bp"),
    ("Risk controls", "Volatility filter, overheat filter, breadth filter, trailing stop"),
]

RESULT_ROWS = [
    ("2023", "Short", "14.92%", "13.58%", "15.15%"),
    ("2023", "Long", "-0.42%", "7.64%", "52.69%"),
    ("2024", "Short", "46.99%", "38.27%", "167.72%"),
    ("2024", "Long", "0.08%", "16.78%", "16.02%"),
    ("2025", "Short", "190.16%", "216.92%", "60.08%"),
    ("2025", "Long", "103.29%", "82.78%", "94.71%"),
    ("2026Q1", "Short", "15.40%", "24.62%", "8.55%"),
    ("2026Q1", "Long", "14.20%", "6.14%", "0.00%"),
]


def load_audit_rows() -> list[tuple[str, str, str, str, str, str]]:
    if not AUDIT_JSON.exists():
        return []
    data = json.loads(AUDIT_JSON.read_text(encoding="utf-8"))
    rows = []
    for row in data.get("rows", []):
        excluded = f"{row.get('future_document_rows_excluded', 0)}/{row.get('future_price_rows_excluded', 0)}"
        rows.append(
            (
                str(row.get("as_of_date", "")),
                str(row.get("document_rows", "")),
                str(row.get("max_document_date", "")),
                str(row.get("price_rows", "")),
                str(row.get("max_price_date", "")),
                excluded,
            )
        )
    return rows


def read_source() -> dict:
    md = SOURCE_MD.read_text(encoding="utf-8")
    title = re.search(r"^#\s+(.+)$", md, re.MULTILINE).group(1)
    english_title = re.search(r"\*\*English Title:\*\*\s*(.+)", md).group(1)
    author = re.search(r"\*\*Author:\*\*\s*(.+)", md).group(1).strip()
    affiliation = re.search(r"\*\*Affiliation:\*\*\s*(.+)", md).group(1).strip()
    english_author_match = re.search(r"\*\*English Author:\*\*\s*(.+)", md)
    english_affiliation_match = re.search(r"\*\*English Affiliation:\*\*\s*(.+)", md)
    english_author = english_author_match.group(1).strip() if english_author_match else author
    english_affiliation = english_affiliation_match.group(1).strip() if english_affiliation_match else affiliation
    abstract = section(md, "ABSTRACT").split("**Key Words:**")[0].strip()
    keywords = re.search(r"\*\*Key Words:\*\*\s*(.+)", md).group(1).strip()
    refs = paragraphs(section(md, "참고문헌").split("## 작성 메모")[0])
    body_sections = []
    for name in ["1. 서론", "2. 제안 프레임워크", "3. 실험 설계 및 결과", "4. 결론"]:
        body_sections.append((name, body_paragraphs(section(md, name))))
    acknowledgements = paragraphs(section(md, "감사의 글"))[0]
    return {
        "title": title,
        "english_title": english_title,
        "author": author,
        "affiliation": affiliation,
        "english_author": english_author,
        "english_affiliation": english_affiliation,
        "abstract": abstract,
        "keywords": keywords,
        "body_sections": body_sections,
        "acknowledgements": acknowledgements,
        "references": refs,
    }


def section(md: str, heading: str) -> str:
    pattern = rf"^##\s+{re.escape(heading)}\s*$"
    match = re.search(pattern, md, re.MULTILINE)
    if not match:
        return ""
    start = match.end()
    next_match = re.search(r"^##\s+", md[start:], re.MULTILINE)
    end = start + next_match.start() if next_match else len(md)
    return md[start:end].strip()


def paragraphs(text: str) -> list[str]:
    return [p.strip().replace("\n", " ") for p in re.split(r"\n\s*\n", text) if p.strip()]


def body_paragraphs(text: str) -> list[str]:
    out: list[str] = []
    for p in paragraphs(text):
        if re.match(r"^Table\s+\d+\.", p) or re.match(r"^Fig\.\s+\d+\.", p) or p.startswith("|"):
            continue
        if p.startswith("`") and p.endswith("`"):
            continue
        if p.startswith("**"):
            continue
        out.append(p)
    return out


def font_path(name: str) -> str | None:
    for candidate in [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts/malgun.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]:
        if candidate.exists():
            return str(candidate)
    return None


def create_architecture_figure(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1300, 430), "white")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(font_path("arialbd.ttf") or font_path("arial.ttf"), 34)
    label_font = ImageFont.truetype(font_path("arial.ttf"), 24)
    small_font = ImageFont.truetype(font_path("arial.ttf"), 19)
    accent = "#2f5f8f"
    fill = "#eef5fb"
    line = "#6f8190"
    draw.text(
        (42, 28),
        "Point-in-Time Temporal Validation Framework",
        fill="#1e2933",
        font=title_font,
    )
    boxes = [
        (60, 125, 250, 255, "Evidence\nSources", "News / DART\nForum / Prices"),
        (310, 125, 500, 255, "Temporal RAG", "as_of_date\nPast-only filter"),
        (560, 125, 770, 255, "Agent Roles", "Analyst / Quant\nChartist / Risk"),
        (830, 125, 1020, 255, "Risk Score", "Volatility / Cost\nGuard filters"),
        (1080, 125, 1250, 255, "Ranking", "Candidate order\nFuture evaluation"),
    ]
    for i, (x1, y1, x2, y2, title, subtitle) in enumerate(boxes):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline=accent, width=3)
        draw.text((x1 + 22, y1 + 18), title, fill=accent, font=label_font)
        draw.multiline_text((x1 + 22, y1 + 64), subtitle, fill="#334155", font=small_font, spacing=5)
        if i < len(boxes) - 1:
            ax1, ay = x2 + 10, (y1 + y2) // 2
            ax2 = boxes[i + 1][0] - 10
            draw.line((ax1, ay, ax2, ay), fill=line, width=4)
            draw.polygon([(ax2, ay), (ax2 - 14, ay - 9), (ax2 - 14, ay + 9)], fill=line)
    draw.text(
        (65, 320),
        "Selection uses only evidence available before each decision date; returns are evaluated after selection.",
        fill="#475569",
        font=small_font,
    )
    img.save(path)


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None, font_name: str = "바탕") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text: str, size_pt: float = 9.0, bold: bool = False, font_name: str = "바탕", align=None):
    paragraph.clear()
    if align is not None:
        paragraph.alignment = align
    add_text_runs(paragraph, text, size_pt=size_pt, bold=bold, font_name=font_name)
    return paragraph


def add_text_runs(paragraph, text: str, size_pt: float = 9.0, bold: bool = False, font_name: str = "바탕") -> None:
    """Add text while rendering numeric reference markers as superscript."""
    parts = re.split(r"(\[\d+(?:[,-]\d+)*\])", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        set_run_font(run, size_pt=size_pt, bold=bold, font_name=font_name)
        if re.fullmatch(r"\[\d+(?:[,-]\d+)*\]", part):
            run.font.superscript = True


def style_exists(doc: Document, name: str) -> bool:
    return name in [style.name for style in doc.styles]


def add_para(doc: Document, text: str = "", style: str | None = "논문 본문", size: float = 9.0, bold: bool = False, align=None):
    p = doc.add_paragraph(style=style if style and style_exists(doc, style) else None)
    set_paragraph_text(p, text, size_pt=size, bold=bold, align=align)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    return p


def add_heading(doc: Document, text: str):
    p = add_para(doc, text, style="각 장 제목", size=10.0, bold=True)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    return p


def clear_cell(cell) -> None:
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)


def add_cell_para(cell, text: str, size=9.0, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, font_name="바탕"):
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold, font_name=font_name)
    return p


def fill_header_table(table, data: dict) -> None:
    for row in table.rows[1:4]:
        for cell in row.cells:
            clear_cell(cell)
    c = table.rows[1].cells[0]
    add_cell_para(c, data["title"], 16, True, font_name="돋움")
    add_cell_para(c, data["english_title"], 13, True, font_name="Times New Roman")
    add_cell_para(c, "", 5)
    add_cell_para(c, f"*,1{data['author']}", 9, False, font_name="돋움")
    add_cell_para(c, f"1{data['affiliation']}", 9, False, font_name="돋움")
    add_cell_para(c, f"*,1{data['english_author']}", 9, False, font_name="Times New Roman")
    add_cell_para(c, f"1{data['english_affiliation']}", 9, False, font_name="Times New Roman")

    c = table.rows[2].cells[0]
    p = add_cell_para(c, "ABSTRACT", 9, True, align=WD_ALIGN_PARAGRAPH.LEFT, font_name="Times New Roman")
    p.paragraph_format.space_after = Pt(2)
    p = add_cell_para(c, data["abstract"], 8.7, False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, font_name="굴림")
    p.paragraph_format.line_spacing = 1.0

    c = table.rows[3].cells[0]
    add_cell_para(
        c,
        f"Key Words: {data['keywords']}",
        8.2,
        False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        font_name="Times New Roman",
    )
    mark_header_row(table.rows[0])


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, size=7.2, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    clear_cell(cell)
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold, font_name="Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_small_table(doc: Document, rows, headers=None):
    n_cols = len(headers) if headers else len(rows[0])
    table = doc.add_table(rows=1 if headers else 0, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    if headers:
        for i, h in enumerate(headers):
            set_cell_text(table.rows[0].cells[i], h, size=6.8, bold=True)
            shade_cell(table.rows[0].cells[i], "E9EEF4")
        mark_header_row(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], str(text), size=6.7, bold=False)
    return table


def mark_section_tables(doc: Document) -> None:
    for section in doc.sections:
        for container in [
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ]:
            for table in container.tables:
                if table.rows:
                    mark_header_row(table.rows[0])


def build_docx() -> None:
    data = read_source()
    create_architecture_figure(FIG_PATH)

    doc = Document(TEMPLATE)
    fill_header_table(doc.tables[0], data)
    mark_section_tables(doc)

    # Keep the title/abstract section break and final two-column section settings.
    body = doc._body._element
    children = list(body)
    for child in children[3:-1]:
        body.remove(child)

    for heading, items in data["body_sections"][:2]:
        add_heading(doc, heading)
        for text in items:
            add_para(doc, text)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(FIG_PATH), width=Inches(3.02))
    for inline in p._p.iter(qn("wp:inline")):
        doc_pr = inline.find(qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", "Architecture diagram of the point-in-time temporal validation framework")
            doc_pr.set("title", "Point-in-time temporal validation framework")
    p.paragraph_format.space_after = Pt(2)
    add_para(
        doc,
        "Fig. 1. Overall process of point-in-time temporal validation.",
        style="Normal",
        size=8.0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    heading, items = data["body_sections"][2]
    add_heading(doc, heading)
    for index, text in enumerate(items):
        add_para(doc, text)
        if index == 1:
            add_para(
                doc,
                "Table 1. Experimental protocol and realism controls.",
                style="Normal",
                size=7.6,
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )
            add_small_table(doc, PROTOCOL_ROWS, headers=["Item", "Setting"])
            add_para(doc, "", style="Normal", size=3)
        if text.startswith("Table 2는"):
            add_para(
                doc,
                "Table 2. Summary of period-horizon comparison.",
                style="Normal",
                size=7.6,
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )
            add_small_table(doc, RESULT_ROWS, headers=["Period", "Horizon", "Hybrid", "Det.", "Best tech"])
            add_para(doc, "", style="Normal", size=3)
        if text.startswith("추가 감사에서는"):
            audit_rows = load_audit_rows()
            if audit_rows:
                add_para(
                    doc,
                    "Table 3. Temporal evidence leakage audit sample.",
                    style="Normal",
                    size=7.6,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                )
                add_small_table(
                    doc,
                    audit_rows,
                    headers=["as_of", "Docs", "Max doc", "Prices", "Max price", "Future d/p"],
                )
                add_para(doc, "", style="Normal", size=3)

    for heading, items in data["body_sections"][3:]:
        add_heading(doc, heading)
        for text in items:
            add_para(doc, text)

    add_heading(doc, "감사의 글")
    add_para(doc, data["acknowledgements"])
    add_heading(doc, "참고문헌")
    for ref in data["references"]:
        add_para(doc, ref, style="참고문헌(내용)", size=7.8)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


def parse_args():
    parser = argparse.ArgumentParser(description="Build the CDE 2026 paper DOCX from Markdown.")
    parser.add_argument("--source-md", default=str(SOURCE_MD), help="Markdown source path")
    parser.add_argument("--out-docx", default=str(OUT_DOCX), help="Output DOCX path")
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    SOURCE_MD = Path(args.source_md)
    OUT_DOCX = Path(args.out_docx)
    build_docx()
